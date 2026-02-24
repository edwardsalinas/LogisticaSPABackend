import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class ThesisDocument:
    def __init__(self, filename="Thesis_Document.docx"):
        self.filename = filename
        self.doc = Document()
        self._setup_margins()
        self._setup_styles()

    def _setup_margins(self):
        """Configura los márgenes según normas APA/ISO (adaptado)."""
        for section in self.doc.sections:
            section.top_margin = Cm(3.03)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(3.0)

    def _setup_styles(self):
        """Configura los estilos de fuente (Arial 12 para texto, TNR para títulos)."""
        # Estilo Normal
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        
        # Estilos de Títulos
        for i in range(1, 4):
            if f'Heading {i}' in self.doc.styles:
                h = self.doc.styles[f'Heading {i}']
                h.font.name = 'Times New Roman'
                h.font.bold = True
                h.font.size = Pt(14 - i)  # H1=13, H2=12, H3=11
                h.font.color.rgb = None  # Color automático (negro)

    def add_heading(self, text, level=1):
        """Agrega un título con el nivel especificado."""
        self.doc.add_heading(text, level)

    def add_paragraph(self, text, bold=False, align=None):
        """Agrega un párrafo. Opcionalmente en negrita o alineado."""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = bold
        if align:
            para.alignment = align
        return para

    def add_bullet_paragraph(self, text, bold=False):
        """Agrega un párrafo con viñeta."""
        para = self.doc.add_paragraph(style='List Bullet')
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = bold
        return para

    def add_image(self, path, width_inches=5.5, caption=None):
        """Agrega una imagen con su leyenda numerada automáticamente."""
        if os.path.exists(path):
            self.doc.add_picture(path, width=Inches(width_inches))
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            print(f"Warning: Image not found at {path}")
            # Placeholder text if image missing
            p = self.doc.add_paragraph(f"[IMAGEN NO ENCONTRADA: {path}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if caption:
            self._add_caption("Figura", caption)

    def add_table_with_data(self, headers, rows, caption=None):
        """Agrega una tabla con datos y leyenda opcional."""
        if caption:
            self._add_caption("Tabla", caption)

        table = self.doc.add_table(rows=1+len(rows), cols=len(headers))
        table.style = 'Table Grid'
        
        # Encabezados
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            # Negrita en headers
            if cell.paragraphs:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Datos
        for ri, row in enumerate(rows):
            for ci, cell_data in enumerate(row):
                table.rows[ri+1].cells[ci].text = str(cell_data)
        
        self.doc.add_paragraph() # Espacio después de la tabla

    def _add_caption(self, label, text):
        """Helper interno para agregar leyendas con numeración automática (SEQ fields)."""
        p_cap = self.doc.add_paragraph(style='Caption')
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # "Figura " o "Tabla "
        run = p_cap.add_run(f"{label} ")
        
        # Campo SEQ para numeración
        self._add_seq_field(run, label)
        
        # Separador y texto
        p_cap.add_run(f": {text}")

    def _add_seq_field(self, run, identifier):
        """Agrega el campo SEQ de Word al run proporcionado."""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = f" SEQ {identifier} \\* ARABIC "
        run._r.append(instr)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar2)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar3)

    def add_toc(self):
        """Agrega una Tabla de Contenidos automática."""
        self._add_toc_field("TOC \\o \"1-3\" \\h \\z \\u")

    def add_index_figures(self):
        """Agrega un Índice de Figuras automático."""
        self._add_toc_field("TOC \\h \\z \\c \"Figura\"")

    def add_index_tables(self):
        """Agrega un Índice de Tablas automático."""
        self._add_toc_field("TOC \\h \\z \\c \"Tabla\"")

    def _add_toc_field(self, instr_text):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = instr_text
        run._r.append(instr)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar2)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar3)

    def add_page_break(self):
        self.doc.add_page_break()

    def save(self):
        self.doc.save(self.filename)
        print(f"Documento guardado: {self.filename}")

    def add_test_case(self, tc_id, name, preconditions, steps, expected, status):
        """Helper para formatear casos de prueba consistentemente."""
        self.add_paragraph(f"Caso de Prueba {tc_id}: {name}", bold=True)
        self.add_paragraph(f"Precondiciones: {preconditions}")
        self.add_paragraph(f"Pasos:\n{steps}")
        self.add_paragraph(f"Resultado esperado: {expected}")
        self.add_paragraph(f"Estado: {status}")
        self.doc.add_paragraph()

if __name__ == "__main__":
    print("Este módulo contiene la clase ThesisDocument para generar documentos.")
